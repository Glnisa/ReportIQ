"""
Word Generator Module for ReportIQ
Creates Word documents with charts and vulnerability analysis
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pathlib import Path
from typing import Optional, Dict, List, Any
from datetime import datetime
from io import BytesIO
import pandas as pd


class WordGenerator:
    """Generates Word documents for vulnerability reports"""
    
    def __init__(self, filter_engine, chart_generator, language: str = "TR"):
        """
        Initialize Word generator
        
        Args:
            filter_engine: FilterEngine instance with filtered data
            chart_generator: ChartGenerator instance with generated charts
            language: Language code (TR or EN)
        """
        self.filter_engine = filter_engine
        self.chart_generator = chart_generator
        self.language = language
        self.document: Optional[Document] = None
        self.vulnerability_dict = None
    
    def set_language(self, language: str) -> None:
        """Set language for document text"""
        self.language = language
    
    def set_vulnerability_dict(self, vuln_dict) -> None:
        """Set vulnerability dictionary for descriptions"""
        self.vulnerability_dict = vuln_dict
    
    def _get_label(self, tr: str, en: str) -> str:
        """Get label based on current language"""
        return tr if self.language == "TR" else en
    
    def _setup_styles(self) -> None:
        """Set up document styles"""
        styles = self.document.styles
        
        # Title style
        title_style = styles['Title']
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 82, 147)
        
        # Heading 1 style
        h1_style = styles['Heading 1']
        h1_style.font.name = 'Arial'
        h1_style.font.size = Pt(18)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 82, 147)
        
        # Heading 2 style
        h2_style = styles['Heading 2']
        h2_style.font.name = 'Arial'
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(51, 51, 51)
        
        # Normal style
        normal_style = styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(11)
    
    def _add_header(self) -> None:
        """Add document header with title and date"""
        # Add logo placeholder or title
        title = self.document.add_paragraph()
        title_run = title.add_run(self._get_label("🛡️ Zafiyet Analiz Raporu", "🛡️ Vulnerability Analysis Report"))
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 82, 147)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = self.document.add_paragraph()
        subtitle_run = subtitle.add_run(self._get_label("Güvenlik Değerlendirme Raporu", "Security Assessment Report"))
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        date_para = self.document.add_paragraph()
        date_text = datetime.now().strftime("%d/%m/%Y %H:%M")
        date_run = date_para.add_run(f"{self._get_label('Rapor Tarihi:', 'Report Date:')} {date_text}")
        date_run.font.size = Pt(11)
        date_run.font.color.rgb = RGBColor(100, 100, 100)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add horizontal line
        self.document.add_paragraph("_" * 80)
        self.document.add_paragraph()
    
    def _add_executive_summary(self) -> None:
        """Add executive summary section"""
        # Section title
        self.document.add_heading(self._get_label("📋 Yönetici Özeti", "📋 Executive Summary"), level=1)
        
        # Get summary statistics
        total = self.filter_engine.get_total_count()
        filtered = self.filter_engine.get_filtered_count()
        sla_summary = self.filter_engine.get_sla_summary()
        priority_summary = self.filter_engine.get_priority_summary()
        
        # Summary paragraph
        summary_text = self._get_label(
            f"Bu rapor, toplam {total:,} kayıttan filtreleme sonucu {filtered:,} zafiyet kaydını analiz etmektedir.",
            f"This report analyzes {filtered:,} vulnerability records filtered from a total of {total:,} records."
        )
        self.document.add_paragraph(summary_text)
        
        # Key metrics table
        self._add_summary_table(filtered, sla_summary, priority_summary)
        
        self.document.add_paragraph()
    
    def _add_summary_table(self, total: int, sla_summary: Dict, priority_summary: Dict) -> None:
        """Add summary metrics table"""
        table = self.document.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Style header row
        header_cells = table.rows[0].cells
        header_cells[0].text = self._get_label("Metrik", "Metric")
        header_cells[1].text = self._get_label("Değer", "Value")
        
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            self._set_cell_shading(cell, "005293")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows
        rows_data = [
            (self._get_label("Toplam Zafiyet", "Total Vulnerabilities"), f"{total:,}"),
            (self._get_label("SLA Dışı", "Out of SLA"), f"{sla_summary.get('out_of_sla', 0):,}"),
            (self._get_label("SLA İçi", "In SLA"), f"{sla_summary.get('in_sla', 0):,}"),
        ]
        
        for i, (metric, value) in enumerate(rows_data, start=1):
            if i < len(table.rows):
                table.rows[i].cells[0].text = metric
                table.rows[i].cells[1].text = value
        
        # Add priority breakdown if space
        priority_text = ", ".join([f"{k}: {v}" for k, v in priority_summary.items()])
        if priority_text:
            row = table.add_row()
            row.cells[0].text = self._get_label("Priority Dağılımı", "Priority Distribution")
            row.cells[1].text = priority_text
    
    def _set_cell_shading(self, cell, color: str) -> None:
        """Set cell background color"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def _add_chart_section(self, title: str, chart_name: str, description: str = "") -> None:
        """Add a chart section to the document"""
        chart_buffer = self.chart_generator.get_chart(chart_name)
        
        if chart_buffer:
            self.document.add_heading(title, level=1)
            
            if description:
                self.document.add_paragraph(description)
            
            # Add chart image
            chart_buffer.seek(0)
            self.document.add_picture(chart_buffer, width=Inches(6.5))
            
            # Center the image
            last_paragraph = self.document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            self.document.add_paragraph()
    
    def _add_top_vulnerabilities_section(self) -> None:
        """Add detailed top vulnerabilities section with descriptions"""
        self.document.add_heading(
            self._get_label("🔥 En Çok Görülen Zafiyetler", "🔥 Most Common Vulnerabilities"), 
            level=1
        )
        
        top_vulns = self.filter_engine.get_top_vulnerabilities(10)
        
        if top_vulns.empty:
            self.document.add_paragraph(
                self._get_label("Zafiyet verisi bulunamadı.", "No vulnerability data found.")
            )
            return
        
        # Add chart if available
        chart_buffer = self.chart_generator.get_chart("top10")
        if chart_buffer:
            chart_buffer.seek(0)
            self.document.add_picture(chart_buffer, width=Inches(6.5))
            last_paragraph = self.document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.document.add_paragraph()
        
        # Add table with details
        self.document.add_heading(
            self._get_label("Zafiyet Detayları", "Vulnerability Details"), 
            level=2
        )
        
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Header row
        headers = [
            self._get_label("Sıra", "Rank"),
            self._get_label("Zafiyet Adı", "Vulnerability Name"),
            self._get_label("Tespit Sayısı", "Count")
        ]
        
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            self._set_cell_shading(cell, "005293")
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows
        for idx, row in top_vulns.iterrows():
            data_row = table.add_row()
            data_row.cells[0].text = str(idx + 1)
            data_row.cells[1].text = str(row['vulnerability'])[:80]
            data_row.cells[2].text = str(row['count'])
        
        self.document.add_paragraph()
        
        # Add descriptions from vulnerability dictionary
        if self.vulnerability_dict and 'plugin_id' in top_vulns.columns:
            self.document.add_heading(
                self._get_label("Zafiyet Açıklamaları", "Vulnerability Descriptions"), 
                level=2
            )
            
            for idx, row in top_vulns.iterrows():
                plugin_id = str(row.get('plugin_id', ''))
                vuln_name = str(row['vulnerability'])
                
                # Try to get description from dictionary
                description = None
                if plugin_id and self.vulnerability_dict:
                    description = self.vulnerability_dict.get_description(
                        plugin_id, 
                        self.language, 
                        ""
                    )
                
                if description:
                    # Add vulnerability name as sub-heading
                    para = self.document.add_paragraph()
                    run = para.add_run(f"• {vuln_name}")
                    run.font.bold = True
                    run.font.size = Pt(11)
                    
                    # Add description
                    desc_para = self.document.add_paragraph(f"  {description}")
                    desc_para.paragraph_format.left_indent = Inches(0.25)
                    
                    # Add remediation if available
                    remediation = self.vulnerability_dict.get_remediation(plugin_id, self.language)
                    if remediation:
                        rem_para = self.document.add_paragraph()
                        rem_run = rem_para.add_run(f"  {self._get_label('Çözüm:', 'Remediation:')} ")
                        rem_run.font.bold = True
                        rem_para.add_run(remediation)
                        rem_para.paragraph_format.left_indent = Inches(0.25)
    
    def _add_department_breakdown(self) -> None:
        """Add department breakdown section"""
        self._add_chart_section(
            self._get_label("🏢 Departman Bazında Analiz", "🏢 Analysis by Department"),
            "department",
            self._get_label(
                "Aşağıdaki grafik, zafiyetlerin departmanlara göre dağılımını göstermektedir.",
                "The chart below shows the distribution of vulnerabilities by department."
            )
        )
        
        # Add department summary table
        dept_data = self.filter_engine.get_count_by_column("department").head(10)
        
        if not dept_data.empty:
            self.document.add_heading(
                self._get_label("Departman Özet Tablosu", "Department Summary Table"), 
                level=2
            )
            
            table = self.document.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            headers = [
                self._get_label("Departman", "Department"),
                self._get_label("Zafiyet Sayısı", "Vulnerability Count")
            ]
            
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                self._set_cell_shading(cell, "005293")
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            
            for dept, count in dept_data.items():
                row = table.add_row()
                row.cells[0].text = str(dept)[:50]
                row.cells[1].text = f"{count:,}"
            
            self.document.add_paragraph()
    
    def _add_filters_applied(self) -> None:
        """Add section showing which filters were applied"""
        self.document.add_heading(
            self._get_label("🔍 Uygulanan Filtreler", "🔍 Applied Filters"), 
            level=2
        )
        
        filters = self.filter_engine.get_filters()
        
        if not filters:
            self.document.add_paragraph(
                self._get_label("Filtre uygulanmadı.", "No filters applied.")
            )
        else:
            for filter_name, value in filters.items():
                if isinstance(value, list):
                    value_str = ", ".join([str(v) for v in value])
                else:
                    value_str = str(value)
                
                para = self.document.add_paragraph()
                run = para.add_run(f"• {filter_name}: ")
                run.font.bold = True
                para.add_run(value_str)
        
        self.document.add_paragraph()
    
    def generate_report(self, sections: List[str], output_path: str) -> str:
        """
        Generate the Word report
        
        Args:
            sections: List of section names to include
            output_path: Path to save the document
            
        Returns:
            Path to the saved document
        """
        self.document = Document()
        self._setup_styles()
        
        # Add header
        self._add_header()
        
        # Add executive summary
        self._add_executive_summary()
        
        # Add filters applied
        self._add_filters_applied()
        
        # Section mappings
        section_methods = {
            "yearly_open": lambda: self._add_chart_section(
                self._get_label("📊 Yıllara Göre Açık Zafiyet", "📊 Open Vulnerabilities by Year"),
                "yearly_open"
            ),
            "priority_dist": lambda: self._add_chart_section(
                self._get_label("🎯 Priority Dağılımı", "🎯 Priority Distribution"),
                "priority_dist"
            ),
            "line_manager": lambda: self._add_chart_section(
                self._get_label("👥 Line Manager Bazında Analiz", "👥 Analysis by Line Manager"),
                "line_manager"
            ),
            "department": self._add_department_breakdown,
            "tool": lambda: self._add_chart_section(
                self._get_label("🔧 Tool/Kaynak Dağılımı", "🔧 Tool/Source Distribution"),
                "tool"
            ),
            "sla": lambda: self._add_chart_section(
                self._get_label("⏰ SLA Durumu", "⏰ SLA Status"),
                "sla"
            ),
            "trend": lambda: self._add_chart_section(
                self._get_label("📈 Trend Analizi", "📈 Trend Analysis"),
                "trend"
            ),
            "top10": self._add_top_vulnerabilities_section,
            "ip_density": lambda: self._add_chart_section(
                self._get_label("💻 IP Bazlı Yoğunluk", "💻 IP-Based Density"),
                "ip_density"
            ),
            "resolution_time": lambda: self._add_chart_section(
                self._get_label("📅 Çözüm Süreleri", "📅 Resolution Times"),
                "resolution_time"
            ),
            "sla_breach": lambda: self._add_chart_section(
                self._get_label("⚠️ SLA Aşım Analizi", "⚠️ SLA Breach Analysis"),
                "sla_breach"
            ),
        }
        
        # Generate requested sections
        for section in sections:
            if section in section_methods:
                try:
                    section_methods[section]()
                except Exception as e:
                    print(f"Error generating section {section}: {e}")
        
        # Add footer
        self._add_footer()
        
        # Save document
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(output))
        
        return str(output)
    
    def _add_footer(self) -> None:
        """Add document footer"""
        self.document.add_paragraph()
        self.document.add_paragraph("_" * 80)
        
        footer = self.document.add_paragraph()
        footer_text = self._get_label(
            f"Bu rapor ReportIQ tarafından {datetime.now().strftime('%d/%m/%Y %H:%M')} tarihinde oluşturulmuştur.",
            f"This report was generated by ReportIQ on {datetime.now().strftime('%Y-%m-%d %H:%M')}."
        )
        run = footer.add_run(footer_text)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
